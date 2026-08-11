import hashlib
from io import BytesIO
import unittest

from docx import Document

from model_doc_agent.docx_exporter import (
    DEFAULT_TEMPLATE_PATH,
    render_markdown_to_docx,
)


class DocxExporterTests(unittest.TestCase):
    def test_bundled_template_matches_supplied_reference(self):
        digest = hashlib.sha256(DEFAULT_TEMPLATE_PATH.read_bytes()).hexdigest()

        self.assertEqual(
            digest,
            "3b3e5ddf863457c45f68651b5c5530b53715ae1293def9101bd49f20f1a828f0",
        )

    def test_markdown_is_rendered_with_template_structure(self):
        markdown = """# Sales — Power BI Documentation

## 1. Document Information

| Field | Details |
|---|---|
| Report Name | Sales |

## 4. Solution Architecture

```mermaid
flowchart LR
    Source --> Model --> Report
```

## 7. Semantic Model

### 7.2 Tables

- **Sales** — fact table (`AI suggestion`)
"""
        output = render_markdown_to_docx(
            markdown,
            document_context={
                "projectName": "Sales Analytics",
                "owner": "BI Team",
                "version": "1.0",
                "dataClassification": "Internal",
            },
        )
        document = Document(BytesIO(output))
        body_text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertTrue(output.startswith(b"PK"))
        self.assertIn("1. Document Information", body_text)
        self.assertIn("4. Solution Architecture", body_text)
        self.assertIn("7. Semantic Model", body_text)
        self.assertIn("Sales Analytics", table_text)
        self.assertIn("Source --> Model --> Report", body_text)
        self.assertNotIn("[Enter details]", body_text + table_text)
        self.assertEqual(
            document.sections[0].header.paragraphs[0].text,
            "POWER BI REPORT DOCUMENTATION",
        )
        self.assertIn(
            "Internal — Power BI technical and functional documentation",
            document.sections[0].footer.paragraphs[0].text,
        )
        self.assertIn(
            "PAGE",
            document.sections[0].footer._element.xml,
        )

    def test_empty_markdown_is_rejected(self):
        with self.assertRaisesRegex(ValueError, "Markdown documentation"):
            render_markdown_to_docx("   ")


if __name__ == "__main__":
    unittest.main()
