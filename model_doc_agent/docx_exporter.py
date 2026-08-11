"""Render agent Markdown as a branded Word document.

The bundled DOCX is the visual authority. Generated documents retain its
theme, styles, page setup, header, footer, and page-number field while the
instructional placeholder body is replaced with the current report.
"""

from io import BytesIO
from pathlib import Path
import re
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE_PATH = (
    PROJECT_ROOT
    / "templates"
    / "Power_BI_Report_Documentation_Template.docx"
)

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "EEF4F8"
WHITE = "FFFFFF"
GRAY = "666666"
MAX_MARKDOWN_LINES = 50_000
MAX_TABLE_COLUMNS = 12


def _clear_body(document: DocumentType) -> None:
    """Remove template placeholder content while retaining section settings."""

    body = document._element.body

    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))

    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)

    shading.set(qn("w:fill"), fill)


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_cell_margins(cell, top: int = 70, start: int = 80,
                      bottom: int = 70, end: int = 80) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")

    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)

    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        margin = margins.find(qn(f"w:{name}"))

        if margin is None:
            margin = OxmlElement(f"w:{name}")
            margins.append(margin)

        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_run_font(run, size: float = 9, color: str = "000000",
                  bold: bool = False, italic: bool = False,
                  name: str = "Aptos") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _add_inline_text(paragraph, value: str, size: float = 9,
                     color: str = "000000", bold: bool = False) -> None:
    """Add the small Markdown subset produced by the documentation agent."""

    token_pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    cursor = 0

    for match in token_pattern.finditer(value):
        if match.start() > cursor:
            run = paragraph.add_run(value[cursor:match.start()])
            _set_run_font(run, size=size, color=color, bold=bold)

        token = match.group(0)

        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            _set_run_font(
                run,
                size=size,
                color=color,
                bold=bold,
                name="Aptos Mono",
            )

        cursor = match.end()

    if cursor < len(value):
        run = paragraph.add_run(value[cursor:])
        _set_run_font(run, size=size, color=color, bold=bold)


def _split_table_row(line: str) -> List[str]:
    value = line.strip()

    if value.startswith("|"):
        value = value[1:]

    if value.endswith("|"):
        value = value[:-1]

    cells = re.split(r"(?<!\\)\|", value)

    return [
        cell.strip().replace(r"\|", "|").replace("<br>", "\n")
        for cell in cells
    ]


def _is_table_separator(cells: Sequence[str]) -> bool:
    return bool(cells) and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))
        for cell in cells
    )


def _normalize_table_rows(rows: Sequence[Sequence[str]]) -> List[List[str]]:
    if not rows:
        return []

    column_count = min(
        max(len(row) for row in rows),
        MAX_TABLE_COLUMNS,
    )
    normalized = []

    for row in rows:
        values = list(row[:column_count])
        values.extend([""] * (column_count - len(values)))
        normalized.append(values)

    return normalized


def _add_branded_table(document: DocumentType,
                       rows: Sequence[Sequence[str]]) -> None:
    normalized_rows = _normalize_table_rows(rows)

    if not normalized_rows:
        return

    table = document.add_table(
        rows=len(normalized_rows),
        cols=len(normalized_rows[0]),
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _repeat_table_header(table.rows[0])

    for row_index, values in enumerate(normalized_rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)

            if row_index == 0:
                _set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, LIGHT_BLUE)

            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.clear()
            _add_inline_text(
                paragraph,
                value,
                size=8,
                color=WHITE if row_index == 0 else "000000",
                bold=row_index == 0,
            )

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_cover_table(document: DocumentType,
                     rows: Sequence[Tuple[str, str]]) -> None:
    table_rows = [["Document Field", "Value"]]
    table_rows.extend([[label, value] for label, value in rows if value])
    _add_branded_table(document, table_rows)


def _document_title(markdown: str) -> str:
    for line in markdown.splitlines():
        match = re.match(r"^#\s+(.+?)\s*$", line)

        if match:
            return match.group(1)

    return "Power BI Report Documentation"


def _short_report_name(title: str) -> str:
    return re.sub(
        r"\s+[—-]\s+Power BI(?: Report)? Documentation\s*$",
        "",
        title,
        flags=re.IGNORECASE,
    ).strip() or title


def _extract_document_fields(markdown: str) -> Dict[str, str]:
    """Read simple Field/Details rows so the cover uses actual report facts."""

    fields = {}

    for line in markdown.splitlines():
        if not line.strip().startswith("|"):
            continue

        cells = _split_table_row(line)

        if len(cells) < 2 or _is_table_separator(cells):
            continue

        key = re.sub(r"\s+", " ", cells[0]).strip().casefold()
        value = cells[1].strip()

        if key in {
            "report name",
            "project name",
            "semantic model",
            "business area",
            "document version",
            "current version",
            "status",
            "last updated",
            "data classification",
        } and value:
            fields.setdefault(key, value)

    return fields


def _add_title_page(document: DocumentType, title: str,
                    context: Dict[str, str],
                    document_fields: Dict[str, str]) -> None:
    power_bi = document.add_paragraph()
    power_bi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    power_bi.paragraph_format.space_before = Pt(85)
    run = power_bi.add_run("POWER BI")
    _set_run_font(run, size=20, color=BLUE, bold=True, name="Aptos Display")

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(title)
    _set_run_font(run, size=26, color=NAVY, bold=True, name="Aptos Display")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run(
        "Technical and functional documentation generated from Power BI "
        "project metadata"
    )
    _set_run_font(run, size=10, color=GRAY, italic=True)

    report_name = (
        document_fields.get("report name")
        or _short_report_name(title)
    )
    project_name = (
        context.get("projectName")
        or document_fields.get("project name", "")
    )

    if project_name.casefold() == report_name.casefold():
        project_name = ""

    cover_rows = [
        ("Report Name", report_name),
        ("Project Name", project_name),
        ("Semantic Model", document_fields.get("semantic model", "")),
        (
            "Business Area",
            context.get("businessArea")
            or document_fields.get("business area", ""),
        ),
        (
            "Document Version",
            context.get("version")
            or document_fields.get("document version")
            or document_fields.get("current version", ""),
        ),
        ("Document Owner", context.get("owner", "")),
        (
            "Status",
            context.get("documentStatus")
            or document_fields.get("status", ""),
        ),
        (
            "Last Updated",
            context.get("revisionDate")
            or document_fields.get("last updated", ""),
        ),
        (
            "Data Classification",
            context.get("dataClassification")
            or document_fields.get("data classification", ""),
        ),
    ]
    _add_cover_table(document, cover_rows)

    note = document.add_paragraph(style="Template Note")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(10)
    note.add_run(
        "Generated from parsed metadata. Organizational details appear only "
        "when supplied by the user."
    )


def _add_heading(document: DocumentType, text: str, level: int,
                 page_break_before: bool = False) -> None:
    # The Markdown H1 is the cover title. H2 therefore maps to Word Heading 1.
    word_level = min(max(level - 1, 1), 3)
    paragraph = document.add_paragraph(style=f"Heading {word_level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.add_run(text)


def _add_code_block(document: DocumentType, lines: Iterable[str],
                    language: str) -> None:
    if language.lower() == "mermaid":
        note = document.add_paragraph(style="Template Note")
        note.add_run(
            "Diagram definition (Mermaid). Paste into a Mermaid-compatible "
            "viewer to render it visually."
        )

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F7F9")
    properties.append(shading)
    borders = OxmlElement("w:pBdr")

    for side in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "D1D5DB")
        borders.append(border)

    properties.append(borders)
    run = paragraph.add_run("\n".join(lines).rstrip())
    _set_run_font(run, size=8, color="1F2937", name="Aptos Mono")


def _is_special_line(line: str) -> bool:
    stripped = line.strip()

    return bool(
        not stripped
        or stripped.startswith("#")
        or stripped.startswith("```")
        or stripped.startswith("|")
        or re.match(r"^[-*]\s+", stripped)
        or re.match(r"^\d+\.\s+", stripped)
    )


def _add_markdown_body(document: DocumentType, markdown: str) -> None:
    lines = markdown.splitlines()

    if len(lines) > MAX_MARKDOWN_LINES:
        raise ValueError("The generated documentation is too large for DOCX export.")

    index = 0
    first_h1_skipped = False
    first_content_heading = True

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        fence = re.match(r"^```\s*([^`]*)$", stripped)

        if fence:
            language = fence.group(1).strip()
            index += 1
            code_lines = []

            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1

            if index < len(lines):
                index += 1

            _add_code_block(document, code_lines, language)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)

        if heading:
            level = len(heading.group(1))
            text = heading.group(2)

            if level == 1 and not first_h1_skipped:
                first_h1_skipped = True
            else:
                _add_heading(
                    document,
                    text,
                    level,
                    page_break_before=first_content_heading,
                )
                first_content_heading = False

            index += 1
            continue

        if stripped.startswith("|"):
            table_rows = []

            while index < len(lines) and lines[index].strip().startswith("|"):
                cells = _split_table_row(lines[index])

                if not _is_table_separator(cells):
                    table_rows.append(cells)

                index += 1

            _add_branded_table(document, table_rows)
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)

        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_text(paragraph, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)

        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_text(paragraph, numbered.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1

        while index < len(lines) and not _is_special_line(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.15
        _add_inline_text(paragraph, " ".join(paragraph_lines))


def _update_package_metadata(document: DocumentType, title: str,
                             context: Dict[str, str]) -> None:
    properties = document.core_properties
    properties.title = title
    properties.subject = "Power BI report technical and functional documentation"
    properties.author = (
        context.get("author")
        or context.get("owner")
        or "Data Model Documentation Agent"
    )
    properties.keywords = "Power BI, semantic model, PBIP, documentation"


def _replace_paragraph_label(paragraph, value: str) -> None:
    """Replace visible label text without disturbing paragraph formatting."""

    if paragraph.runs:
        paragraph.runs[0].text = value

        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _update_headers_and_footers(document: DocumentType,
                                context: Dict[str, str]) -> None:
    """Adapt template labels for a completed report and keep the PAGE field."""

    classification = context.get("dataClassification", "")
    footer_label = "Power BI technical and functional documentation"

    if classification:
        footer_label = f"{classification} — {footer_label}"

    for section in document.sections:
        if section.header.paragraphs:
            _replace_paragraph_label(
                section.header.paragraphs[0],
                "POWER BI REPORT DOCUMENTATION",
            )

        if section.footer.paragraphs:
            _replace_paragraph_label(
                section.footer.paragraphs[0],
                footer_label,
            )


def render_markdown_to_docx(
    markdown: str,
    document_context: Optional[Dict[str, str]] = None,
    template_path: Optional[Path] = None,
) -> bytes:
    """Return a DOCX byte string using the bundled organizational template."""

    if not isinstance(markdown, str) or not markdown.strip():
        raise ValueError("Markdown documentation is required for DOCX export.")

    selected_template = template_path or DEFAULT_TEMPLATE_PATH

    if not selected_template.exists():
        raise RuntimeError(
            "The Word documentation template is missing from the project."
        )

    context = {
        key: value.strip()
        for key, value in (document_context or {}).items()
        if isinstance(value, str) and value.strip()
    }
    document = Document(str(selected_template))
    _clear_body(document)

    title = _document_title(markdown)
    document_fields = _extract_document_fields(markdown)
    _add_title_page(document, title, context, document_fields)
    _add_markdown_body(document, markdown)
    _update_package_metadata(document, title, context)
    _update_headers_and_footers(document, context)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
